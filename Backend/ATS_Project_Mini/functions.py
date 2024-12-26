import re
import json
from PyPDF2 import PdfReader
from docx import Document
import spacy
import os
from django.conf import settings

# Load NLP model
nlp = spacy.load("en_core_web_sm")

def load_job_criteria(json_path):
    with open(json_path, "r") as f:
        return json.load(f)

def extract_text_from_pdf(filepath):
    print("File path received by extract text resume in pdf")
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# def extract_text_from_docx(filepath):
#     print("File path received by extract text resume in docs")
#     doc = Document(filepath)
#     print("Docs converted-----> ", doc)
#     return "\n".join([p.text for p in doc.paragraphs])
def extract_text_from_docx(filepath):
    doc = Document(filepath)
    # Include text from paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Handle additional content like tables
    tables = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    tables.append(cell_text)
    # Combine paragraphs and table text
    return "\n".join(paragraphs + tables)

def socialCheck(text, json_path=os.path.join(settings.BASE_DIR, 'ATS_Project_Mini', 'static', 'job_criteria.json')):
    """
    Function to check social contact details in the resume.
    """
    # Load social patterns from the JSON file
    def load_social_patterns(json_path):
        print("Json reaches to Socail ",json_path)
        with open(json_path, "r") as f:
            criteria = json.load(f)
        return criteria.get("social_patterns", {})

    social_patterns = load_social_patterns(json_path)

    # Initialize results dictionary
    results = {
        "email": [],
        "phone": [],
        "linkedin": [],
        "github": [],
        "missing_email": False,
        "missing_phone": False,
        "missing_linkedin": False,
        "missing_github": False,
    }

    # Check for matches using the loaded patterns
    for key, pattern in social_patterns.items():
        matches = re.findall(pattern, text)
        results[key] = matches
        if not matches:
            results[f"missing_{key}"] = True
    print("Socail Checking done and found this: ",results)
    return results

def analyze_resume(filepath, job_criteria):
    if filepath.endswith(".pdf"):
        print("File path received by analyze resume in pdf")
        text = extract_text_from_pdf(filepath)
    elif filepath.endswith(".docx"):
        print("File path received by analyze resume in docs")
        text = extract_text_from_docx(filepath)
    else:
        raise ValueError("Unsupported file format!")

    doc = nlp(text)
    results = {
        "keyword_matches": [],
        "missing_keywords": [],
        "skills_matches": [],
        "missing_skills": [],
        "education_matches": [],
        "missing_education": [],
        "format_issues": [],
    }

    resume_keywords = [token.text.lower() for token in doc]
    for keyword in job_criteria["keywords"]:
        if keyword.lower() in resume_keywords:
            results["keyword_matches"].append(keyword)
        else:
            results["missing_keywords"].append(keyword)

     # Skill matching with scoring logic
    skills_matched = [skill for skill in job_criteria["skills"] if skill.lower() in resume_keywords]
    results["skills_matches"] = skills_matched
    skill_count = len(skills_matched)
    results["skill_score"] = round(skill_count * 0.2, 2)  # Each skill contributes 0.2 points

    if skill_count < 5:
        results["missing_skills"] = [f"Include at least 5 skills. Currently, {skill_count} skill(s) matched."]


    resume_text = text.lower()





    # Code for cheking the presence of education degree in the resume # 


        # for edu in job_criteria["education"]:
    #     if edu.lower() in resume_text:
    #         results["education_matches"].append(edu)
    #     else:
    #         results["missing_education"].append(edu)



    # results = {
    #     "education_matches": [],
    #     "missing_education": [],
    # }

    bachelors_matched = []
    for bachelor in job_criteria["education"]["bachelors"]:
        if bachelor.lower() in resume_text:
            bachelors_matched.append(bachelor)

    # if bachelors_matched :
    #     masters_matched = []
    #     for master in job_criteria["education"]["masters"]:
    #         if master.lower() in resume_text:
    #             masters_matched.append(master)


    # Handle bachelor's degree matches
    if bachelors_matched:
        results["education_matches"].append("Bachelors: " + ", ".join(bachelors_matched))
    else:
        results["missing_education"].append("No Bachelor's degree found.")

    #  # Handle master's degree matches
    # if masters_matched:
    #     results["education_matches"].append("Masters: " + ", ".join(masters_matched))

    # # Handle missing education
    # if not bachelors_matched and not masters_matched:
    #     results["missing_education"].append("No education information found.")




    # Education cheking function ends here

    if "table" in resume_text or "image" in resume_text:
        results["format_issues"].append("Contains tables or images")

    # Call socialCheck with the correct argument name
    social_results = socialCheck(text, json_path=os.path.join(settings.BASE_DIR, 'ATS_Project_Mini', 'static', 'job_criteria.json'))
    results.update(social_results)

    return results, text

def calculate_score(results, job_criteria):
    total_criteria = len(job_criteria["keywords"]) + len(job_criteria["education"]) + 1  # Add 1 for skills section
    matched_criteria = len(results["keyword_matches"]) + len(results["education_matches"])
    matched_criteria += results["skill_score"]  # Include the skill score

    for key in ["email", "phone", "linkedin", "github"]:
        if not results.get(f"missing_{key}", True):
            matched_criteria += 1

    score = (matched_criteria / total_criteria) * 100
    return round(score, 2)

def generate_feedback(results, output_path=None):
    feedback = []
    if results["missing_keywords"]:
        feedback.append(f"Missing Keywords: {', '.join(results['missing_keywords'])}")
    if results["missing_skills"]:
        feedback.append(f"Missing Skills: {' '.join(results['missing_skills'])}")
    if results["missing_education"]:
        feedback.append(f"Missing Education: {', '.join(results['missing_education'])}")
    if results["format_issues"]:
        feedback.append(f"Formatting Issues: {', '.join(results['format_issues'])}")

    for key in ["email", "phone", "linkedin", "github"]:
        if results.get(f"missing_{key}", True):
            feedback.append(f"Missing {key.capitalize()}: Include your {key} in the resume.")

    # Default output path in MEDIA_ROOT
    if not output_path:
        output_path = os.path.join(settings.MEDIA_ROOT, "feedback.txt")

    with open(output_path, "w") as f:
        f.write("\n".join(feedback))
    return output_path











def generate_appreciation(results, output_path="appreciation.txt"):
    appreciation = []

    # Check for matched keywords
    if results["keyword_matches"]:
        appreciation.append(f"Keywords Matched: {', '.join(str(keyword) for keyword in results['keyword_matches'])}")
    
    # Check for matched skills
    if results["skills_matches"]:
        appreciation.append(f"Skills Matched: {', '.join(str(skill) for skill in results['skills_matches'])}")
    
    # Check for matched education
    if results["education_matches"]:
        appreciation.append(f"Education Matches: {', '.join(str(edu) for edu in results['education_matches'])}")
    
    # Check for matched social details (email, phone, linkedin, github)
    social_matches = []
    for key in ["email", "phone", "linkedin", "github"]:
        if results.get(f"{key}", []):
            social_matches.append(f"{key.capitalize()}: {', '.join(str(item) for item in results[key])}")

    if social_matches:
        appreciation.append(f"Social Details Matched: {', '.join(social_matches)}")
    
    # If there are any matched items, write them to the appreciation file
    if appreciation:
        with open(output_path, "w") as f:
            f.write("\n".join(appreciation))
        print(f"Appreciation file generated at: {output_path}")
    else:
        print("No matched points found to appreciate.")
